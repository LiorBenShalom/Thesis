#!/usr/bin/env python3
"""
Extract and Tag Citations from Legal Verdicts

- Identifies citations using regex patterns (ע"פ, רע"פ, etc.)
- Filters by relevant sections (sentencing range, discussion, etc.)
- Extracts context using GPT
- Tags relevance using BERT (heBERT)
"""

import re
import os
import gc
import pandas as pd
from pathlib import Path
from tqdm import tqdm
import argparse

# Try to import optional dependencies
try:
    import torch
    from transformers import BertTokenizer, BertForSequenceClassification
    HAS_TORCH = True
except ImportError:
    HAS_TORCH = False
    print("⚠️ torch/transformers not available - skipping BERT tagging")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx not available")

try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False
    print("⚠️ openai not available - skipping GPT extraction")

# ========== LEGAL ACRONYMS ==========
acronyms = [
    "אב", "אבע", "אימוצ", "אמצ", "אפ", "אפח", "את", "אתפ", "באפ", "באש", "בבנ", "בגצ", "בדא", "בדמ",
    "בדמש", "בהנ", "בהע", "בהש", "בידמ", "בידע", "בל", "בלמ", "במ", "בעא", "בעח", "בעמ", "בעק", "בפ",
    "בפמ", "בפת", "בצא", "בצהמ", "בק", "בקמ", "בקשה", "ברמ", "ברע", "ברש", "בש", "בשא",
    "בשגצ", "בשהת", "בשז", "בשמ", "בשע", "בשפ", "בתת", "גזז", "גמר", "גפ", "דבע", "דח", "דט", "דיונ",
    "דמ", "דמר", "דמש", "דנ", "דנא", "דנגצ", "דנמ", "דנפ", "הד", "הדפ", "הוצלפ", "הט", "הכ", "המ",
    "המד", "הממ", "המע", "המש", "הנ", "הסת", "הע", "העז", "הפ", "הפב", "הפמ", "הצמ", "הש", "השא",
    "השגצ", "השפ", "השר", "הת", "וחק", "וע", "ושמ", "ושק", "ושר", "זי", "חא", "חבר", "חד", "חדא",
    "חדלפ", "חדלת", "חדמ", "חדפ", "חהע", "חי", "חנ", "חסמ", "חעמ", "חעק", "חש", "יוש", "ייתא", "ימא",
    "יס", "כצ", "מ", "מא", "מבכ", "מבס", "מונופולינ", "מזג", "מח", "מחוז", "מחע", "מט", "מטכל", "מי",
    "מיב", "מכ", "ממ", "מס", "מסט", "מעי", "מעת", "מקמ", "מרכז", "מת", "נ", "נב", "נבא", "נמ", "נמב",
    "נעד", "נער", "סבא", "סע", "סעש", "סק", "סקכ", "ע", "עא", "עאח", "עאפ", "עב", "עבאפ", "עבז", "עבח",
    "עבי", "עבל", "עבמצ", "עבעח", "עבפ", "עבר", "עבשהת", "עגר", "עדי", "עדמ", "עהג", "עהס", "עהפ",
    "עו", "עורפ", "עז", "עח", "עחא", "עחדלפ", "עחדפ", "עחדת", "עחהס", "עחע", "עחק", "עחר", "עכב",
    "על", "עלא", "עלבש", "עלח", "עלע", "עמ", "עמא", "עמה", "עמז", "עמח", "עמי", "עמלע", "עממ", "עמנ",
    "עמפ", "עמצ", "עמק", "עמרמ", "עמש", "עמשמ", "עמת", "ענ", "ענא", "ענמ", "ענמא", "ענמש", "ענפ",
    "עסא", "עסק", "עע", "עעא", "עעמ", "עער", "עעתא", "עפ", "עפא", "עפג", "עפהג", "עפמ", "עפמק",
    "עפנ", "עפס", "עפספ", "עפע", "עפר", "עפת", "עצמ", "עק", "עקג", "עקמ", "עקנ", "עקפ", "ער", "ערא",
    "ערגצ", "ערמ", "ערעור", "ערפ", "ערר", "עש", "עשא", "עשמ", "עשר", "עשת", "עשתש", "עת", "עתא",
    "עתמ", "עתפב", "עתצ", "פא", "פה", "פל", "פלא", "פמ", "פמר", "פעמ", "פקח", "פר", "פרק", "פשז",
    "פשר", "פת", "צא", "צבנ", "צה", "צו", "צח", "צמ", "קג", "קפ", "רחדפ", "רמש", "רע", "רעא", "רעב",
    "רעבס", "רעו", "רעמ", "רעס", "רעפ", "רעפא", "רעצ", "רער", "רערצ", "רעש", "רעתא", "רצפ", "רתק",
    "ש", "שבד", "שמ", "שמי", "שנא", "שע", "שעמ", "שק", "שש", "תא", "תאדמ", "תאח", "תאמ", "תאק", "תב",
    "תבכ", "תבע", "תג", "תגא", "תד", "תדא", "תהג", "תהנ", "תהס", "תוב", "תוח", "תח", "תחפ", "תחת",
    "תט", "תי", "תכ", "תלא", "תלב", "תלהמ", "תלפ", "תלתמ", "תמ", "תמהח", "תממ", "תמק", "תמר",
    "תמש", "תנג", "תנז", "תע", "תעא", "תעז", "תפ", "תפב", "תפח", "תפחע", "תפכ", "תפמ", "תפע",
    "תפק", "תצ", "תק", "תקח", "תקמ", "תרמ", "תת", "תתח", "תתע", "תתעא", "תתק"
]

# Required sections for filtering
required_parts = [
    "מתחמי ענישה", "אחידות בענישה", "מתחם הענישה", "מתחם ענישה", "דיון",
    "ענישה נהוגה", "הענישה הנוהגת", "ענישה נוהגת", "מתחם העונש", "מתחם עונש",
    "מדיניות הענישה", "והכרעה", "ההרשעה", "מדיניות הענישה הנהוגה"
]


def create_acronym_variants(acronyms):
    """Create regex pattern for all acronym variants."""
    acronym_variants = []
    for a in acronyms:
        if len(a) > 1:
            base_acronym = a
            if a.startswith('ב') or a.startswith('ו') or a.startswith('ה'):
                base_acronym = a[1:]
            
            for acr in [a, base_acronym]:
                if len(acr) > 1:
                    quoted = rf"{acr[:-1]}[\"'״]{acr[-1]}"
                    with_dot = rf"{acr[:-1]}\.{acr[-1]}"
                    acronym_variants.append(f"(?:{quoted}|{with_dot})")
                    dots_between = r'\.'.join(list(acr))
                    acronym_variants.append(dots_between)
    
    return '|'.join(acronym_variants)


def clean_leading_prefix(citation):
    """Clean leading prefix letters from citation."""
    match = re.match(r'^([לבוה])\s*([א-ת"]+)', citation)
    if not match:
        return citation
    prefix = match.group(1)
    maybe_acronym = match.group(2)
    
    full = prefix + maybe_acronym
    
    def normalize(text):
        return text.replace('"', '').replace("״", "").replace("'", "").replace("׳", "")
    
    norm_maybe = normalize(maybe_acronym)
    norm_full = normalize(full)
    
    if norm_maybe in acronyms and norm_full not in acronyms:
        return citation[len(prefix):].lstrip()
    
    return citation


def normalize_case_name(name):
    """Normalize case name for comparison."""
    if pd.isna(name):
        return ""
    name = str(name)
    name = re.sub(r"\(.*?\)", "", name)
    name = re.sub(r"[∕/\\]", "-", name)
    name = re.sub(r"\s+", " ", name)
    name = name.strip().lower().replace(" ", "_")
    return name


# Build citation regex
acronym_pattern = create_acronym_variants(acronyms)
number_pattern = r'''
    (?:
        \d{1,6}[-/]\d{2}[-/]\d{2}
        | \d{1,6}[-/]\d{1,6}
        | \d{1,6}-\d{2}-\d{2}
    )
'''
citation_pattern = fr'''
    (?<!\w)
    ([א-ת]?)
    ({acronym_pattern})
    \.?
    \s*
    (\((.*?)\))?
    \s*[-/]?\s*
    ({number_pattern})
    (?!\w)
'''.strip()

citation_regex = re.compile(citation_pattern, re.VERBOSE)


def extract_citations_from_csv(csv_data):
    """Extract all citations from CSV data."""
    citations = []
    text_column = csv_data["text"].astype(str)
    matches = text_column.str.extractall(citation_regex)
    
    for _, row in matches.iterrows():
        citation = " ".join(map(str, filter(pd.notna, row))).strip()
        citation = re.sub(r"\s{2,}", " ", citation)
        
        if re.match(r"^על \d+$", citation):
            continue
        
        citation = re.sub(r"\((.*?)\)\s+\1", r"(\1)", citation)
        citation = clean_leading_prefix(citation)
        citations.append(citation)
    
    return citations if citations else []


def filter_csv_relevant_parts(csv_data):
    """Filter CSV to relevant parts only. If no relevant parts found, return full CSV."""
    start_index = None
    
    for idx, row in csv_data.iterrows():
        if any(req_part in str(row.get("part", "")) for req_part in required_parts):
            start_index = idx
            break
    
    if start_index is not None:
        return csv_data.iloc[start_index:]
    else:
        # No relevant parts found - return full CSV instead of skipping
        return csv_data


def find_all_occurrences(doc, citation):
    """Find all paragraph indices containing the citation."""
    indices = []
    for i, paragraph in enumerate(doc.paragraphs):
        if citation in paragraph.text:
            indices.append(i)
    return indices


def get_context_paragraphs(doc, index, citation):
    """Get context paragraphs around a citation."""
    context_text = []
    
    prev_index = index - 1
    while prev_index >= 0 and not doc.paragraphs[prev_index].text.strip():
        prev_index -= 1
    
    if prev_index >= 0:
        context_text.append(doc.paragraphs[prev_index].text.strip())
    
    curr_text = doc.paragraphs[index].text.strip()
    if curr_text:
        context_text.append(curr_text)
    else:
        return None
    
    next_index = index + 1
    while next_index < len(doc.paragraphs) and not doc.paragraphs[next_index].text.strip():
        next_index += 1
    
    if next_index < len(doc.paragraphs):
        context_text.append(doc.paragraphs[next_index].text.strip())
    
    if not context_text:
        return None
    
    return "\n".join(context_text).strip()


def query_gpt_batch_extraction(citation_contexts, client):
    """Use a single GPT call to extract relevant text for ALL citations in a document.
    
    Args:
        citation_contexts: list of (citation, context_text) tuples
        client: OpenAI client
    
    Returns:
        dict mapping citation -> extracted_text
    """
    numbered_items = []
    for i, (citation, context) in enumerate(citation_contexts, 1):
        numbered_items.append(
            f"[{i}] Citation: {citation}\nText:\n{context}"
        )
    
    all_items = "\n\n---\n\n".join(numbered_items)
    
    prompt = f"""Given the following legal texts, each with a citation, extract **only** the part of the text that directly relates to each citation.

**Rules:**
- Do not modify any wording. Keep original phrasing.
- Do not summarize or rephrase.
- Return only the relevant portion for each citation.
- If a citation appears in a list, include the preceding explanation.

For each item, respond with:
[number] extracted text here

{all_items}"""

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are an AI trained to extract legal citations. Respond with numbered extractions matching the input numbers."},
                {"role": "user", "content": prompt}
            ]
        )
        raw = response.choices[0].message.content
        
        extractions = {}
        current_num = None
        current_lines = []
        
        for line in raw.split('\n'):
            bracket_match = re.match(r'^\[(\d+)\]\s*(.*)', line)
            if bracket_match:
                if current_num is not None:
                    extractions[current_num] = '\n'.join(current_lines).strip()
                current_num = int(bracket_match.group(1))
                current_lines = [bracket_match.group(2)]
            elif current_num is not None:
                current_lines.append(line)
        
        if current_num is not None:
            extractions[current_num] = '\n'.join(current_lines).strip()
        
        result = {}
        for i, (citation, context) in enumerate(citation_contexts, 1):
            result[citation] = extractions.get(i, context)
        
        return result
    except Exception as e:
        print(f"🚨 GPT API error: {e}")
        return {citation: context for citation, context in citation_contexts}


def process_document(docx_path, csv_path, output_path, model_bert=None, tokenizer_bert=None, client=None, device="cpu"):
    """Process a single document and extract/tag citations."""
    if not HAS_DOCX:
        print("⚠️ python-docx required")
        return
    
    doc = docx.Document(docx_path)
    csv_data = pd.read_csv(csv_path)
    filtered_csv_data = filter_csv_relevant_parts(csv_data)
    
    if filtered_csv_data.empty:
        return
    
    citations = extract_citations_from_csv(filtered_csv_data)
    
    if len(citations) > 30:
        print(f"⚠️ Too many citations ({len(citations)}) in {docx_path}")
        return
    
    print(f"🔍 Found {len(citations)} citations")
    
    citation_contexts = []
    for citation in citations:
        citation_indices = find_all_occurrences(doc, citation)
        
        merged_contexts = []
        for index in citation_indices:
            full_context = get_context_paragraphs(doc, index, citation)
            if full_context:
                merged_contexts.append(full_context)
        
        if not merged_contexts:
            continue
        
        final_context = "\n".join(set(merged_contexts)).strip()
        citation_contexts.append((citation, final_context))
    
    if not citation_contexts:
        return
    
    if HAS_OPENAI and client:
        extracted_map = query_gpt_batch_extraction(citation_contexts, client)
    else:
        extracted_map = {cit: ctx for cit, ctx in citation_contexts}
    
    results = []
    for citation, final_context in citation_contexts:
        extracted_text = extracted_map.get(citation, final_context)
        
        prediction = 0
        if HAS_TORCH and model_bert and tokenizer_bert:
            encoding = tokenizer_bert(extracted_text, truncation=True, padding=True, max_length=128, return_tensors="pt")
            encoding = {key: val.to(device) for key, val in encoding.items()}
            with torch.no_grad():
                output = model_bert(**encoding)
                prediction = torch.argmax(output.logits, dim=-1).item()
        
        result = {
            'citation': normalize_case_name(citation),
            'context_text': final_context,
            'extracted_text': extracted_text,
            'predicted_label': prediction
        }
        results.append(result)
    
    if results:
        results_df = pd.DataFrame(results)
        results_df.to_csv(output_path, index=False, encoding="utf-8")
        print(f"✅ Saved {len(results)} citations to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="חילוץ ותיוג ציטוטים מפסקי דין")
    parser.add_argument("--docx-dir", required=True, help="תיקיית קבצי DOCX")
    parser.add_argument("--csv-dir", required=True, help="תיקיית קבצי CSV")
    parser.add_argument("--output-dir", required=True, help="תיקיית פלט")
    parser.add_argument("--model-path", help="נתיב למודל BERT (אופציונלי)")
    
    args = parser.parse_args()
    
    docx_directory = Path(args.docx_dir)
    csv_directory = Path(args.csv_dir)
    output_directory = Path(args.output_dir)
    output_directory.mkdir(parents=True, exist_ok=True)
    
    # Initialize models if available
    device = "cuda" if HAS_TORCH and torch.cuda.is_available() else "cpu"
    model_bert = None
    tokenizer_bert = None
    client = None
    
    if HAS_TORCH and args.model_path:
        model_path = Path(args.model_path)
        if model_path.is_dir() and (model_path / "config.json").exists():
            tokenizer_bert = BertTokenizer.from_pretrained(str(model_path))
            model_bert = BertForSequenceClassification.from_pretrained(str(model_path))
            print(f"✅ Loaded BERT model from directory {model_path}")
        elif model_path.is_file():
            tokenizer_bert = BertTokenizer.from_pretrained('avichr/heBERT')
            model_bert = BertForSequenceClassification.from_pretrained('avichr/heBERT', num_labels=2)
            model_bert.load_state_dict(torch.load(args.model_path, map_location=device))
            print(f"✅ Loaded BERT model from {args.model_path}")
        else:
            print(f"⚠️ Model path not found: {args.model_path}; running without BERT")
            model_bert = None
            tokenizer_bert = None
        if model_bert is not None:
            model_bert.to(device)
            model_bert.eval()
    
    if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        print("✅ OpenAI client initialized")
    
    # Process files
    all_files = list(docx_directory.glob("*.docx"))
    print(f"🗂 Found {len(all_files)} DOCX files")
    
    for file_path in tqdm(all_files, desc="Processing"):
        csv_file = csv_directory / f"{file_path.stem}.csv"
        output_file = output_directory / f"{file_path.stem}.csv"
        
        if not csv_file.exists():
            continue
        
        if output_file.exists() and output_file.stat().st_size > 0:
            continue
        
        process_document(
            str(file_path), str(csv_file), str(output_file),
            model_bert, tokenizer_bert, client, device
        )
        gc.collect()
    
    print("✅ Done!")


if __name__ == "__main__":
    main()



